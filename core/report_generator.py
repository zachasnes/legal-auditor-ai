from docx import Document
from io import BytesIO

def create_word_report(target_name: str, review_mode: str, client_type: str, combined_markdown_text: str) -> BytesIO:
    """Takes combined LLM markdown text, parses the table, and returns a BytesIO docx buffer."""
    
    doc = Document()
    doc.add_heading(f'Legal Audit Report: {target_name}', 0)
    doc.add_paragraph(f"Mode: {review_mode}\nClient: {client_type}")

    # Professional 3-column table layout
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Clause Reference'
    hdr_cells[1].text = 'Our Standard'
    hdr_cells[2].text = 'Their Substantive Change'

    # Attempt to parse markdown table from combined responses
    lines = combined_markdown_text.strip().split('\n')
    table_lines = [line.strip() for line in lines if line.strip().startswith('|')]
    
    if len(table_lines) > 2: # At least header, separator, and one data row
        # skip header and separator (index 0 and 1)
        for row_line in table_lines[2:]:
            # split by | and remove empty strings from ends
            cols = [c.strip() for c in row_line.split('|')][1:-1]
            if len(cols) >= 3:
                row_cells = table.add_row().cells
                row_cells[0].text = cols[0]
                row_cells[1].text = cols[1]
                row_cells[2].text = cols[2]
            elif len(cols) > 0:
                row_cells = table.add_row().cells
                row_cells[0].text = cols[0]
                row_cells[1].text = "..." if len(cols) < 2 else cols[1]
                row_cells[2].text = "..." if len(cols) < 3 else cols[2]
    else:
        # Fallback if the AI returned plain text instead of a table
        row_cells = table.add_row().cells
        row_cells[0].text = "Full Document Review"
        row_cells[1].text = "See next column"
        row_cells[2].text = combined_markdown_text

    # Save to buffer
    bio = BytesIO()
    doc.save(bio)
    return bio
