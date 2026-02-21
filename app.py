import streamlit as st
import google.generativeai as genai
import os
import PyPDF2
from docx import Document
from io import BytesIO

# --- PAGE CONFIG ---
st.set_page_config(page_title="Legal Auditor: Professional Edition", page_icon="⚖️", layout="wide")
# --- CUSTOM CSS FOR SPLIT VIEW ---
st.markdown("""
<style>
    .reportview-container { background: #f8f9fa; }
    .stHeader { color: #002D62; }
    .comparison-card { background: #ffffff; padding: 15px; border-radius: 8px; border-left: 5px solid #002D62; margin-bottom: 10px; }
</style>
""", unsafe_allow_html=True)

# --- AUTHENTICATION ---
if "auth" not in st.session_state:
    st.session_state.auth = False

SECURE_PASSWORD = os.environ.get("APP_PASSWORD")

if not st.session_state.auth:
    st.title("🔒 Attorney Access Only")
    entered_password = st.text_input("Access Password", type="password")
    if entered_password == SECURE_PASSWORD:
        st.session_state.auth = True
        st.rerun()
    st.stop()

# --- ENGINE SETUP (AUTO-DISCOVERY) ---
api_key = os.environ.get("GOOGLE_API_KEY")
genai.configure(api_key=api_key)

@st.cache_resource
def get_best_model():
    try:
        available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        preferences = ['models/gemini-1.5-pro', 'models/gemini-1.5-flash', 'models/gemini-pro']
        for p in preferences:
            if p in available_models:
                return genai.GenerativeModel(p), p
        return genai.GenerativeModel(available_models[0]), available_models[0]
    except:
        return None, "Error Connecting"

model, model_name = get_best_model()

# --- HEADER ---
st.title("⚖️ Legal Auditor: Professional Edition")
st.caption(f"Engine: {model_name} | Mode: High-Precision Comparison")

# --- UI LAYOUT: TWO COLUMNS ---
col_left, col_right = st.columns(2)

with col_left:
    st.header("1. Your Source of Truth")
    st.info("Upload your **Playbook** or **Your Last Sent Draft**.")
    source_files = st.file_uploader("Upload Source PDF(s)", type="pdf", accept_multiple_files=True, key="source")
    
    source_text = ""
    if source_files:
        for f in source_files:
            reader = PyPDF2.PdfReader(f)
            for p in reader.pages:
                source_text += f"\n<source_doc name='{f.name}'>\n{p.extract_text()}\n</source_doc>"

with col_right:
    st.header("2. The Counterparty Draft")
    st.warning("Upload the **Incoming Document** you need to review.")
    target_files = st.file_uploader("Upload Review PDF(s)", type="pdf", accept_multiple_files=True, key="target")

# --- SETTINGS & TRIGGER ---
st.divider()
c1, c2, c3 = st.columns([1, 1, 1])

with c1:
    client_type = st.selectbox("I represent the:", ["Landlord", "Tenant", "Buyer", "Seller", "Lender", "Borrower"])

with c2:
    review_mode = st.radio(
        "Review Logic:",
        ["**Playbook Match** (Concept Search)", "**Literal Redline** (Word-for-Word)"],
        help="Playbook Match looks for legal concepts. Literal Redline flags every word change."
    )

if st.button("🚀 Run Comparison Audit"):
    if not source_text or not target_files:
        st.error("Missing Files: Please upload both a 'Source' and a 'Review' document.")
    else:
        for target in target_files:
            st.subheader(f"🔍 Deviation Report: {target.name}")
            
            with st.spinner("Processing..."):
                target_reader = PyPDF2.PdfReader(target)
                target_text = ""
                for p in target_reader.pages:
                    target_text += p.extract_text() + "\n"

                # --- PROFESSIONAL PROMPT (NOISE-REDUCED) ---
                # Prioritizing substantive legal changes over formatting/typos
                prompt_task = (
                    "STRICT SUBSTANTIVE COMPARISON. Ignore all whitespace, punctuation, "
                    "typos, or formatting changes. Only flag changes to legal meaning."
                    if "Literal" in review_mode else 
                    "CONCEPTUAL PLAYBOOK COMPLIANCE. Flag deviations from mandatory standards."
                )
                
                prompt = f"""
                You are a senior attorney's high-precision comparison tool. 
                Representing: {client_type}
                Task: {prompt_task}

                [SOURCE OF TRUTH / PLAYBOOK]
                {source_text}

                [INCOMING DRAFT FOR REVIEW]
                {target_text}

                CRITICAL INSTRUCTIONS:
                1. IGNORE ALL TYPOS AND SPACING: Do not flag 't he' vs 'the' or missing spaces.
                2. IGNORE PUNCTUATION: Do not flag commas or periods unless they change liability.
                3. FOCUS: Only list changes that alter the legal rights or obligations of the {client_type}.
                4. FORMAT: Create a Markdown table with 3 columns:
                   - Clause Reference
                   - Our Standard
                   - Their Substantive Change
                5. Do NOT provide summaries. If no substantive changes exist, return "No substantive deviations found."
                """

                try:
                    # Token limit set to 8k to prevent cutoff at page 11
                    response = model.generate_content(
                        prompt,
                        generation_config={
                            "max_output_tokens": 8192,
                            "temperature": 0.1
                        }
                    )
                    st.markdown(response.text)
                    
                    # --- SIDE-BY-SIDE WORD REPORT ---
                    doc = Document()
                    doc.add_heading(f'Legal Audit Report: {target.name}', 0)
                    doc.add_paragraph(f"Mode: {review_mode}\nClient: {client_type}")

                    # Professional 2-column table layout
                    table = doc.add_table(rows=1, cols=2)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Original Segment'
                    hdr_cells[1].text = 'Substantive Deviation & Analysis'

                    row_cells = table.add_row().cells
                    row_cells[0].text = "Review of Full Document"
                    row_cells[1].text = response.text

                    bio = BytesIO()
                    doc.save(bio)
                    st.download_button(
                        label="💾 Download Professional Report",
                        data=bio.getvalue(),
                        file_name=f"Audit_{target.name}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                except Exception as e:
                    st.error(f"AI Engine Error: {str(e)}")